import json
import os
from docx import Document
import re
import string
import requests
import time
import datetime
#import textract
import codecs
import platform
import pythoncom
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from io import StringIO

EXTENSION_DOC = 'doc'
EXTENSION_DOCX = 'docx'
EXTENSION_PDF = 'pdf'
EXTENSION_TXT = 'txt'

print(f"platform = {platform.platform().lower()}")


def fix_continuous_paragraph(content):
    '''
    :param content: content of file, original text - not lower
    :return: text with modified_paragraph
    '''
    list_result = []
    list_para = content.split('\n')
    i = 0
    while i < len(list_para):
        current_para = list_para[i]
        if i == len(list_para) - 1:
            list_result.append(current_para)
            break

        j = i + 1
        # remain_list = list_para[j:]
        while j < len(list_para):
            next_para = list_para[j]
            if next_para is None or len(next_para) == 0:
                j += 1
            elif next_para[0].islower():
                current_para = f'{current_para} {next_para}'
                j += 1
            else:
                break
        list_result.append(current_para)
        i = j
    return '\n'.join(list_result)


def fix_each_pdf_paragraph(content):
    content = content.replace('', '')
    content = content.replace('\n', ' ')
    content = re.sub('[ ]+', ' ', content)
    return content


def remove_number_page(content_in):
    list_result = []
    list_para_split = content_in.split('\n')
    index_skip = [-1]
    # if len(list_para_split) < 4:
    #     return list
    for i in range(len(list_para_split) - 1):
        if i in index_skip:
            continue
        para = list_para_split[i].strip()
        next_para = list_para_split[i + 1].strip()
        if re.fullmatch(r"[0-9]+", next_para):
            if i == len(list_para_split) - 2:
                list_result.append(para)
                break
            else:
                third_para = list_para_split[i + 2].strip()
                if third_para == '':
                    list_result.append(para)
                    index_skip = [i + 1, i + 2]
                    continue
                first_token = third_para[0]
                if first_token.islower():
                    list_result.append(f"{para} {third_para}")
                    index_skip = [i + 1, i + 2]
                else:
                    list_result.append(para)
                    index_skip = [i + 1]
        else:
            list_result.append(para)
            if i == len(list_para_split) - 2:
                list_result.append(next_para)
                break
            # list_result.append(next_para)
            # index_skip = [i + 1]
    return '\n'.join(list_result)


def read_doc_file(path_file):
    """
    :param path_file:
    :return: None if error, else return read content
    """
    platform_os = platform.platform().lower()
    result = None
    if 'win' in platform_os:
        # read in windows
        # import aspose.words as aw
        # doc = aw.Document(path_file)
        # path_docx_file = path_file + 'x'
        # doc.save(path_docx_file)
        # if not os.path.exists(path_docx_file):
        #     return result
        # result = read_docx_file(path_docx_file)
        # result = result.replace("Evaluation Only. Created with Aspose.Words. Copyright 2003-2022 Aspose Pty Ltd.", "")
        # result = result.replace(
        #     "Created with an evaluation copy of Aspose.Words. To discover the full versions of our APIs please visit: https://products.aspose.com/words/",
        #     "")
        # result = result.replace("This document was truncated here because it was created in the Evaluation Mode.", "")

        import win32com.client as win32
        docx_file_path = os.path.splitext(os.path.abspath(path_file))[0] + ".docx"
        if not os.path.exists(docx_file_path):
            # Initialize
            pythoncom.CoInitialize()

            #word = win32.gencache.EnsureDispatch('Word.Application')
            word = dispatch('Word.Application')
            doc = word.Documents.Open(path_file)
            doc.Activate()
            new_file_abs = os.path.splitext(os.path.abspath(docx_file_path))[0] + ".docx"
            word.ActiveDocument.SaveAs(
                new_file_abs, FileFormat=win32.constants.wdFormatXMLDocument
            )
            doc.Close(False)

            if not os.path.exists(docx_file_path):
                return result
            result = read_docx_file(docx_file_path)

            # os.system('antiword ' + path_file + ' > ' + docx_file)
            # with open(docx_file) as f:
            #     result = f.read()
            # os.remove(docx_file)  # docx_file was just to read, so deleting
            #
            # list_paras = result.split('\n\n')
            # result = '\n'.join(fix_each_pdf_paragraph(para) for para in list_paras)
            # result = remove_number_page(result)
        else:
            # already a file with same name as doc exists having docx extension,
            # which means it is a different file, so we cant read it
            print('Info : file with same name of doc exists having docx extension, so we cant read it')
            result = None
    elif 'mac' in platform_os:
        # read in macOS
        #text = textract.process(path_file)
        #text = codecs.decode(text, encoding='utf-8')
        #list_paras = text.split('\n\n')
        #result = '\n'.join(fix_each_pdf_paragraph(para) for para in list_paras)
        result = remove_number_page(result)
    return result


def read_docx_file(path_file):
    document = Document(path_file)
    paras = document.paragraphs
    result = '\n'.join([para.text for para in paras])
    return result


def dispatch(app_name:str):
    try:
        from win32com import client
        app = client.gencache.EnsureDispatch(app_name)
    except AttributeError:
        # Corner case dependencies.
        import os
        import re
        import sys
        import shutil
        # Remove cache and try again.
        MODULE_LIST = [m.__name__ for m in sys.modules.values()]
        for module in MODULE_LIST:
            if re.match(r'win32com\.gen_py\..+', module):
                del sys.modules[module]
        shutil.rmtree(os.path.join(os.environ.get('LOCALAPPDATA'), 'Temp', 'gen_py'))
        from win32com import client
        app = client.gencache.EnsureDispatch(app_name)
    return app

# def read_pdf_file(path_file):
#     text = textract.process(path_file, method='pdfminer')
#     text = codecs.decode(text, encoding='utf-8')
#     list_paras = text.split('\n\n')
#     result = '\n'.join(fix_each_pdf_paragraph(para) for para in list_paras)
#     result = remove_number_page(result)
#     return result
def recognise_topic_sentence(text_in):
    list_para = text_in.split('\n')
    list_result = []
    for para in list_para:
        if re.fullmatch(r"^[IXV1-9.]+\.\s.*$", para):
            list_result.append(f"{para}\n")
        else:
            list_result.append(para)
    return '\n'.join(list_result)


def read_pdf_file(path_pdf):
    rsrcmgr = PDFResourceManager()
    retstr = StringIO()
    laparams = LAParams()
    device = TextConverter(rsrcmgr, retstr, codec='utf-8', laparams=laparams)

    fp = open(path_pdf, 'rb')
    pagenos = set()
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    for page in PDFPage.get_pages(fp, pagenos, maxpages=0, password="", caching=True,
                                  check_extractable=True):
        interpreter.process_page(page)
    text = retstr.getvalue()

    fp.close()
    device.close()
    retstr.close()

    list_paras = text.split('\n\n')
    text_temp = '\n\n'.join([recognise_topic_sentence(x) for x in list_paras])
    list_paras = text_temp.split('\n\n')
    result = '\n'.join(fix_each_pdf_paragraph(para) for para in list_paras)
    result = remove_number_page(result)
    return result


def extract_text_from_file(path_file):
    '''
    :param path_file: path of input file (doc, docx, pdf)
    :return: content of text file or None if error
    '''
    extension = get_file_name_extension(path_file)
    result = None
    message = ''
    try:
        if extension == EXTENSION_DOCX:
            result = read_docx_file(path_file)
        elif extension == EXTENSION_DOC:
            result = read_doc_file(path_file)
        elif extension == EXTENSION_PDF:
            result = read_pdf_file(path_file)
        if result:
            # fix paragraphs
            result = fix_continuous_paragraph(result)
        if result == '':
            result = None
            message = 'File is empty!'
    except Exception as e:
        result = None
        message = str(e)
        print(f"Read file at path: {path_file} fail with exception: {message}")
    return result, message


def get_session_id(config):
    return datetime.datetime.now().strftime(config['format_session_id'])


def get_current_date():
    return datetime.datetime.now().strftime("%d-%m-%Y")


def get_config():
    current_path = os.getcwd()
    # current_path = os.path.dirname(os.path.abspath(__file__))
    config_link = os.path.join(current_path, 'config_run.json')
    f = open(config_link)
    data = json.load(f)
    config_data = data['config']
    f.close()
    return config_data


def get_token(config):
    username = config['username_get_token']
    password = config['password_get_token']
    url = config['api_get_token']
    headers = {'Content-Type': 'application/x-www-form-urlencoded'}
    r = requests.get(url, data={'Username': username, 'Password': password, 'grant_type': 'password'}, headers=headers)
    data = r.json()
    access_token = data['access_token']
    r.close()
    return access_token


def read_all_files(root_folder):
    return [os.path.join(root_folder, i) for i in os.listdir(root_folder) if
            os.path.isfile(os.path.join(root_folder, i))]


def remove_number_words(content, number_skipping_words):
    number_words = len(content.split())
    if number_skipping_words >= number_words - 20:
        return None
    temp = content.replace('\n', ' \n ')
    all_words = re.split('[ ]+', temp)
    # all_words = temp.split(r'[ ]+')
    all_words_copy = all_words.copy()
    number_removed_words = 0

    for word in all_words:
        if number_removed_words == number_skipping_words:
            break
        if word != '\n':
            all_words_copy.remove(word)
            number_removed_words += 1
    result = ' '.join(all_words_copy)
    return result.strip()


def make_dir(directory):
    if not os.path.isdir(directory):
        os.mkdir(directory)
    return directory


def read_all_path_file(root_dir):
    # return list of path files in root_dir directory
    return [os.path.join(root_dir, file_) for file_ in os.listdir(root_dir) if
            os.path.isfile(os.path.join(root_dir, file_))]


def get_file_name_extension(path_file):
    return os.path.splitext(path_file)[1].replace('.', '')


def filter_file_extension(list_files, list_extension_keep=None):
    # filter files from given extension
    # list_extension_keep = [f'.{ext}' for ext in list_extension_keep]
    if list_extension_keep is None:
        list_extension_keep = [EXTENSION_DOC, EXTENSION_DOCX, EXTENSION_PDF]
    list_files_result = []
    for file in list_files:
        # extension = os.path.splitext(file)[1]
        extension = get_file_name_extension(file)
        if extension in list_extension_keep:
            list_files_result.append(file)
    return list_files_result


def get_current_time_milliseconds():
    return round(time.time() * 1000)


def save_text_result(text_result, path_file, dir_out):
    make_dir(dir_out)
    basename = os.path.basename(path_file).split(".")[0]
    output_path = os.path.join(dir_out, f'{basename}.txt')
    f = open(output_path, 'w')
    f.write(text_result)
    f.close()


def format_file_name(path_file_name):
    path_root_dir = os.path.dirname(path_file_name)
    filename = os.path.basename(path_file_name)

    # need to be modified
    file_name_without_ext = os.path.splitext(filename)[0]

    extension = os.path.splitext(filename)[1]

    file_name_without_ext = re.sub(u'[àáạảãâầấậẩẫăằắặẳẵ]', 'a', file_name_without_ext)
    file_name_without_ext = re.sub(u'[ÀÁẠẢÃĂẰẮẶẲẴÂẦẤẬẨẪ]', 'A', file_name_without_ext)
    file_name_without_ext = re.sub(u'[èéẹẻẽêềếệểễ]', 'e', file_name_without_ext)
    file_name_without_ext = re.sub(u'[ÈÉẸẺẼÊỀẾỆỂỄ]', 'E', file_name_without_ext)
    file_name_without_ext = re.sub(u'[òóọỏõôồốộổỗơờớợởỡ]', 'o', file_name_without_ext)
    file_name_without_ext = re.sub(u'[ÒÓỌỎÕÔỒỐỘỔỖƠỜỚỢỞỠ]', 'O', file_name_without_ext)
    file_name_without_ext = re.sub(u'[ìíịỉĩ]', 'i', file_name_without_ext)
    file_name_without_ext = re.sub(u'[ÌÍỊỈĨ]', 'I', file_name_without_ext)
    file_name_without_ext = re.sub(u'[ùúụủũưừứựửữ]', 'u', file_name_without_ext)
    file_name_without_ext = re.sub(u'[ƯỪỨỰỬỮÙÚỤỦŨ]', 'U', file_name_without_ext)
    file_name_without_ext = re.sub(u'[ỳýỵỷỹ]', 'y', file_name_without_ext)
    file_name_without_ext = re.sub(u'[ỲÝỴỶỸ]', 'Y', file_name_without_ext)
    file_name_without_ext = re.sub(u'[ĐÐ]', 'D', file_name_without_ext)
    file_name_without_ext = re.sub(u'[đ]', 'd', file_name_without_ext)
    file_name_without_ext = file_name_without_ext.lower().strip()
    # punc
    file_name_without_ext = re.sub(r'[%s-–]' % (string.punctuation.replace("_", "")), '_', file_name_without_ext)
    file_name_without_ext = re.sub(r'\s+', '_', file_name_without_ext).strip()
    if (convert_int(file_name_without_ext) > 0):
        file_name_without_ext = ("c_%s" % file_name_without_ext)
    return f'{path_root_dir}/{file_name_without_ext}{extension}'


def convert_float(value):
    try:
        result = float(value)
        return result
    except BaseException as e:
        return 0.0


def convert_int(value):
    try:
        result = int(round(convert_float(value)))
        return result
    except BaseException as e:
        return 0


def read_tables_docx_file(path_file):
    document = Document(path_file)
    result = document.tables
    return result



'''
def read_tables_pdf_file(path_pdf):
    result = camelot.read_pdf(path_pdf,pages='all')
    return result
'''

def extract_table_from_file(path_file):
    '''
    :param path_file: path of input file (doc, docx, pdf)
    :return: content of text file or None if error
    '''
    extension = get_file_name_extension(path_file)
    result = None
    message = ''
    try:
        if extension == EXTENSION_DOCX:
            result = read_tables_docx_file(path_file)
        #elif extension == EXTENSION_DOC:
            #result = read_doc_file(path_file)
        elif extension == EXTENSION_PDF:
            result = read_tables_pdf_file(path_file)
        if result == '':
            result = None
            message = 'File is empty!'
    except Exception as e:
        result = None
        message = str(e)
        print(f"Read file at path: {path_file} fail with exception: {message}")
    return result, message

# if __name__ == '__main__':
#     x = extract_text_from_file(
#         '/System/Volumes/Data/DataTrungAnh/VIETTEL_PROJECT/BoNN/code/ReportAnalysis/test_file/pdf_chuan.pdf')
#     f = open('./oooo.txt', 'w')
#     f.write(x)
#     f.close()
#     print(x)
