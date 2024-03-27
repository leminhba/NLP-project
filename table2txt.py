from docx import Document
import pandas as pd

# Hàm chuyển đổi bảng trong Document thành văn bản
def read_table(table):
    # Đọc dữ liệu từ bảng
    data = []
    for row in table.rows:
        row_data = [cell.text for cell in row.cells]
        data.append(row_data)

    # Chuyển dữ liệu thành DataFrame
    df = pd.DataFrame(data[1:], columns=data[0])
    return df

# Hàm chuyển đổi DataFrame thành văn bản
def convert_df_to_text(df):
    text_list = []
    for index, row in df.iterrows():
        text = ", ".join([f"{col}: {row[col]}" for col in df.columns])
        text_list.append(text)
    return "\n".join(text_list)


# Hàm để đọc và xử lý toàn bộ nội dung file .docx
def process_docx(docx_path, output_txt_path):
    doc = Document(docx_path)
    full_text = []

    for para in doc.paragraphs:
        full_text.append(para.text)

    for table in doc.tables:
        table_text = convert_df_to_text(read_table(table))
        full_text.append(table_text)

    with open(output_txt_path, 'w', encoding='utf-8') as f:
        for text in full_text:
            f.write(text + "\n")

# Đường dẫn đến tài liệu .docx và file .txt đầu ra
docx_path = 'D:/3. DM ketqua ban hanh VB theo Ke hoach 30.12.docx'
output_txt_path = 'D:/output_document.txt'

# Xử lý và ghi ra file
process_docx(docx_path, output_txt_path)
